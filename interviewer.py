import os
import time
import logging
import re
import threading
import queue
import tempfile
import json
from docx import Document
from langchain.vectorstores import FAISS
from langchain.embeddings import OllamaEmbeddings
from langchain.llms import Ollama
from langchain.schema import Document as LC_Document
import speech_recognition as sr
import pyttsx3
import sounddevice as sd
import soundfile as sf
import numpy as np
import wave
from pydub import AudioSegment
from pydub.playback import play
from typing import List, Tuple, Optional, Dict

# Setup logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# Model configurations
MODELS = {
    "primary": "deepseek-r1:latest",      # 4.7GB - Heavy lifting
    "secondary": "gemma3:4b",           # 3.3GB - Quick responses
    "embedding": "nomic-embed-text:latest" # 274MB - Efficient embeddings
}

# Initialize models
llm_primary = Ollama(model=MODELS["primary"])
llm_secondary = Ollama(model=MODELS["secondary"])
embedding_model = OllamaEmbeddings(model=MODELS["embedding"])

class TTSManager:
    def __init__(self):
        self.engine = self._init_tts_engine()
        
    def _init_tts_engine(self):
        engine = pyttsx3.init()
        engine.setProperty('rate', 150)
        engine.setProperty('volume', 0.9)
        voices = engine.getProperty('voices')
        if voices:
            female_voices = [v for v in voices if 'female' in v.name.lower()]
            engine.setProperty('voice', female_voices[0].id if female_voices else voices[0].id)
        return engine
    
    def speak(self, text: str):
        print(f"🔊 Speaking: {text}")
        self.engine.say(text)
        self.engine.runAndWait()

class STTManager:
    def __init__(self):
        self.recognizer = sr.Recognizer()
        self.recognizer.energy_threshold = 300
        self.recognizer.dynamic_energy_threshold = True
        self.recognizer.pause_threshold = 2

    def listen(self, timeout: Optional[int] = None, prompt: Optional[str] = None) -> str:
        if prompt:
            print(prompt)
            tts_manager.speak(prompt)
        
        print("🎤 Listening... (speak now)")
        audio_queue = queue.Queue()
        stop_event = threading.Event()
        
        def audio_callback(indata, frames, time, status):
            if status:
                print(f"Audio status: {status}")
            audio_queue.put(indata.copy())
        
        with sd.InputStream(callback=audio_callback, channels=1, samplerate=16000):
            recorded_frames = []
            silence_start = None


            
            recording_start = time.time()
            
            while not stop_event.is_set():
                try:
                    data = audio_queue.get(timeout=0.5)
                    recorded_frames.append(data)
                    if np.abs(data).mean() < 0.01:
                        silence_start = silence_start or time.time()
                        if silence_start and time.time() - silence_start > 1.5:
                            stop_event.set()
                    else:
                        silence_start = None
                    if timeout and time.time() - recording_start > timeout:
                        stop_event.set()
                except queue.Empty:
                    continue
        
        if not recorded_frames:
            print("No audio detected.")
            return ""
        
        audio_data = np.concatenate(recorded_frames, axis=0)
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_file:
            with wave.open(temp_file.name, 'wb') as wf:
                wf.setnchannels(1)
                wf.setsampwidth(2)
                wf.setframerate(16000)
                wf.writeframes((audio_data * 32767).astype(np.int16).tobytes())
            
            with sr.AudioFile(temp_file.name) as source:
                audio = self.recognizer.record(source)
        
        try:
            text = self.recognizer.recognize_google(audio)
            print(f"🎤 Recognized: {text}")
            return text
        except sr.UnknownValueError:
            print("Speech was not understood")
            return ""
        except sr.RequestError as e:
            print(f"Could not request results; {e}")
            return ""
        finally:
            os.remove(temp_file.name)

# Global TTS and STT instances
tts_manager = TTSManager()
stt_manager = STTManager()

class InterviewAssistant:
    def __init__(self, use_speech: bool = True):
        self.use_speech = use_speech
        self.vectorstore = None

    def get_input(self, prompt: str) -> str:
        if self.use_speech:
            print(f"\n{prompt} (speak or type 't:' followed by text)")
            tts_manager.speak(prompt)
            user_choice = input("Press Enter to speak, or type 't:' followed by text: ").strip()
            return stt_manager.listen(timeout=120) if not user_choice.startswith("t:") else user_choice[2:].strip()
        return input(f"{prompt}: ").strip()

    def load_resume(self, file_path: str) -> str:
        try:
            if file_path.endswith(".docx"):
                doc = Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs]).strip()
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read().strip()
        except Exception as e:
            logging.error(f"Error loading resume: {e}")
            return ""

    def setup_vectorstore(self, resume_text: str) -> FAISS:
        faiss_path = "resume_memory"
        try:
            if os.path.exists(faiss_path):
                vectorstore = FAISS.load_local(faiss_path, embedding_model, allow_dangerous_deserialization=True)
                existing_docs = vectorstore.similarity_search(resume_text[:100], k=1)
                if not existing_docs or existing_docs[0].page_content != resume_text:
                    vectorstore.add_documents([LC_Document(page_content=resume_text)])
                    vectorstore.save_local(faiss_path)
            else:
                vectorstore = FAISS.from_documents([LC_Document(page_content=resume_text)], embedding_model)
                vectorstore.save_local(faiss_path)
            return vectorstore
        except Exception as e:
            logging.error(f"Error with vectorstore: {e}")
            return FAISS.from_documents([LC_Document(page_content=resume_text)], embedding_model)

    def generate_questions(self, resume_text: str, self_intro: str, num_questions: int = 5) -> List[str]:
        retriever = self.vectorstore.as_retriever()
        context = "\n".join([doc.page_content for doc in retriever.invoke("Generate interview questions")])
        combined_context = f"Self-introduction: {self_intro}\n\nResume: {context}"
        
        prompt = f"""Generate {num_questions} specific interview questions based on:
{combined_context}
Rules:
1. Relate to skills/experience mentioned
2. Include one technical question about strongest skill
3. Include one problem-solving question
4. Include one follow-up from self-introduction
5. Format as numbered list (1., 2., etc.)"""
        
        try:
            response = llm_primary.invoke(prompt)
            questions = [re.sub(r"^\d+[\.\)\-] *", "", q.strip()) for q in response.split("\n") if "?" in q]
            return questions[:num_questions] if len(questions) >= num_questions else questions + [
                "Can you describe a challenging project you worked on?" for _ in range(num_questions - len(questions))
            ]
        except Exception as e:
            logging.error(f"Error generating questions: {e}")
            return ["Tell me about your technical background?"] * num_questions

    def generate_followup(self, question: str, answer: str, resume_text: str, asked: List[str]) -> str:
        prompt = f"""Generate ONE follow-up question based on:
Question: {question}
Answer: {answer}
Resume: {resume_text[:500]}...
Asked: {asked}
Rules: Dig deeper, be specific, avoid repeats"""
        try:
            followup = llm_secondary.invoke(prompt).strip()
            return re.sub(r"^(Follow-up|Question):", "", followup).strip() + ("?" if "?" not in followup else "")
        except Exception:
            return "Can you provide more details on that?"

    def evaluate_response(self, question: str, answer: str, resume: str, intro: str) -> Tuple[float, str]:
        prompt = f"""Evaluate this answer (1-10):
Question: {question}
Answer: {answer}
Resume: {resume[:500]}...
Intro: {intro[:500]}...
Criteria: Relevance (30%), Depth (30%), Clarity (20%), Context (20%)
Format: [SCORE]: explanation"""
        try:
            eval = llm_secondary.invoke(prompt)
            score = float(re.search(r"(\d+(?:\.\d+)?)", eval).group(1))
            explanation = eval.replace(str(score), "").strip("[]:")
            return score, explanation
        except Exception:
            return 5.0, "Evaluation error"

    def conduct_interview(self, questions: List[str], resume: str, intro: str) -> Tuple[List[float], List[str], List[str], List[str]]:
        scores, feedback, all_questions, answers = [], [], questions.copy(), []
        
        for i, q in enumerate(questions, 1):
            print(f"\nQ{i}: {q}")
            tts_manager.speak(q)
            answer = self.get_input("Your answer") or ""
            answers.append(answer)
            
            if not answer:
                scores.append(0)
                feedback.append("No answer provided")
                continue
            
            score, exp = self.evaluate_response(q, answer, resume, intro)
            scores.append(score)
            feedback.append(exp)
            print(f"Score: {score:.1f}/10 - {exp}")
            tts_manager.speak(f"Score: {int(score)}/10. {exp}")
            
            if i < len(questions) and score < 7.0:
                followup = self.generate_followup(q, answer, resume, all_questions)
                if followup not in all_questions:
                    all_questions.append(followup)
                    print(f"\nFollow-up: {followup}")
                    tts_manager.speak(followup)
                    f_answer = self.get_input("Your answer") or ""
                    answers.append(f_answer)
                    if f_answer:
                        f_score, f_exp = self.evaluate_response(followup, f_answer, resume, intro)
                        scores.append(f_score)
                        feedback.append(f_exp)
                        print(f"Score: {f_score:.1f}/10 - {f_exp}")
        
        return scores, feedback, all_questions, answers

    def final_assessment(self, scores: List[float], feedback: List[str], questions: List[str], 
                        answers: List[str], resume: str, intro: str) -> Tuple[float, str]:
        avg_score = sum(scores) / len(scores) if scores else 0
        qa_summary = "".join(
            f"Q{i}: {q}\nA: {a}\nScore: {s:.1f}/10 - {f}\n\n" 
            for i, (q, a, s, f) in enumerate(zip(questions, answers, scores, feedback), 1)
        )
        
        prompt = f"""Assess this interview:
{qa_summary}
Resume: {resume[:500]}...
Intro: {intro[:500]}...
Score: {avg_score:.1f}/10
Include: 3 strengths, 3 improvements, 2 tips"""
        
        try:
            return avg_score, llm_primary.invoke(prompt)
        except Exception:
            return avg_score, f"Score: {avg_score:.1f}/10. Assessment generation failed."

    def save_results(self, score: float, assessment: str, questions: List[str], answers: List[str], 
                    scores: List[float], feedback: List[str]) -> str:
        os.makedirs("interview_results", exist_ok=True)
        timestamp = time.strftime("%Y%m%d-%H%M%S")
        filename = f"interview_results/interview_{timestamp}.txt"
        
        with open(filename, "w") as f:
            f.write(f"INTERVIEW RESULTS - {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write(f"Overall Score: {score:.1f}/10\n\n")
            for i, (q, a, s, fb) in enumerate(zip(questions, answers, scores, feedback), 1):
                f.write(f"Q{i}: {q}\nA: {a}\nScore: {s:.1f}/10 - {fb}\n\n")
            f.write("ASSESSMENT:\n" + assessment)
        
        print(f"\n📝 Saved to '{filename}'")
        tts_manager.speak("Results saved to file.")
        return filename

def main():
    print("🤖 Interactive Interview Assistant")
    tts_manager.speak("Welcome to the Interactive Interview Assistant!")
    
    use_speech = input("Use speech recognition? (y/n): ").lower().startswith('y')
    assistant = InterviewAssistant(use_speech)
    
    resume_path = assistant.get_input("📂 Enter resume file path")
    if not os.path.exists(resume_path):
        print("❌ Resume not found!")
        return
    
    resume_text = assistant.load_resume(resume_path) or "No details provided."
    assistant.vectorstore = assistant.setup_vectorstore(resume_text)
    
    intro = assistant.get_input("🎤 Please introduce yourself (2 min)") or "No introduction provided."
    
    questions = assistant.generate_questions(resume_text, intro)
    print("\nGenerated Questions:")
    for i, q in enumerate(questions, 1):
        print(f"Q{i}: {q}")
    
    if assistant.get_input("Start interview? (y/n)").lower().startswith('y'):
        scores, feedback, all_questions, answers = assistant.conduct_interview(questions, resume_text, intro)
        final_score, assessment = assistant.final_assessment(scores, feedback, all_questions, answers, resume_text, intro)
        
        print(f"\n🏆 FINAL ASSESSMENT\nScore: {final_score:.1f}/10\n{assessment}")
        tts_manager.speak(f"Final score: {int(final_score)}/10.")
        assistant.save_results(final_score, assessment, all_questions, answers, scores, feedback)
    else:
        print("Interview canceled.")
        tts_manager.speak("Interview canceled.")

if __name__ == "__main__":
    main()
